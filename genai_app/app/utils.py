import os
import pytesseract
from PIL import Image
from PyPDF2 import PdfReader
import docx
import pandas as pd
import sqlite3
import whisper
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# ----------------------------------------
# Configuration
# ----------------------------------------
# Point to the actual tesseract.exe file, not just the folder
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'

# Embedding setup
embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
base_vectorstore_path = os.path.join("..", "vector_store")

# ----------------------------------------
# File extraction logic
# ----------------------------------------
def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower().strip('.')
    try:
        if ext == "pdf":
            return _extract_pdf(path)
        elif ext == "txt":
            return _extract_txt(path)
        elif ext == "docx":
            return _extract_docx(path)
        elif ext in ["xlsx", "csv"]:
            return _extract_spreadsheet(path)
        elif ext == "db":
            return extract_sqlite(path)
        elif ext in ["jpg", "jpeg", "png"]:
            return _extract_image(path)
        elif ext in ["mp4", "mp3", "wav"]:
            return transcribe_media(path)
    except Exception as e:
        print(f"⚠ Error extracting .{ext} file at '{path}': {e}")
    return ""

def _extract_pdf(path: str) -> str:
    reader = PdfReader(path)
    return " ".join(page.extract_text() or "" for page in reader.pages)

def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def _extract_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join(para.text for para in doc.paragraphs)

def _extract_spreadsheet(path: str) -> str:
    # Supports both .xlsx and .csv
    if path.lower().endswith(".csv"):
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    return df.to_csv(index=False)

def _extract_image(path: str) -> str:
    # Open in binary mode to reduce permission issues
    with Image.open(path) as img:
        # Optional: convert to RGB if you have weird modes
        img = img.convert("RGB")
        return pytesseract.image_to_string(img)

# ----------------------------------------
# SQLite table text extraction
# ----------------------------------------
def extract_sqlite(path: str) -> str:
    conn = sqlite3.connect(path)
    text = ""
    try:
        tables = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table';"
        ).fetchall()
        for (table_name,) in tables:
            df = pd.read_sql_query(f"SELECT * FROM \"{table_name}\";", conn)
            text += df.to_csv(index=False)
    except Exception as e:
        print(f"⚠ SQLite extract error for '{path}': {e}")
    finally:
        conn.close()
    return text

# ----------------------------------------
# Transcribe media (audio/video)
# ----------------------------------------
def transcribe_media(path: str) -> str:
    model = whisper.load_model("base")
    result = model.transcribe(path)
    return result.get("text", "")

# ----------------------------------------
# Embed and store with LangChain + Chroma
# ----------------------------------------
def embed_and_store(text: str, user_id: str) -> None:
    chunks = split_text(text)
    documents = [Document(page_content=chunk) for chunk in chunks]
    vector_path = os.path.join(base_vectorstore_path, user_id)

    # Ensure directory exists
    os.makedirs(vector_path, exist_ok=True)
    # Optional: clear old files
    for fname in os.listdir(vector_path):
        os.remove(os.path.join(vector_path, fname))

    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embedding_model,
        persist_directory=vector_path
    )
    vectordb.persist()
    print(f"[Chroma] Stored {len(documents)} chunks for user '{user_id}'")

# ----------------------------------------
# Retrieve relevant chunks for query
# ----------------------------------------
def retrieve_relevant_chunks(user_id: str, query: str, k: int = 3) -> str:
    vector_path = os.path.join(base_vectorstore_path, user_id)
    try:
        vectordb = Chroma(
            persist_directory=vector_path,
            embedding_function=embedding_model
        )
        retriever = vectordb.as_retriever(search_kwargs={"k": k})
        docs = retriever.get_relevant_documents(query)
        return "\n".join(doc.page_content for doc in docs)
    except Exception as e:
        print(f"[Chroma Retrieval Error for user '{user_id}']: {e}")
        return ""

# ----------------------------------------
# Chunking logic
# ----------------------------------------
def split_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    return splitter.split_text(text)
